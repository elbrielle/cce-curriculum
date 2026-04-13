#!/usr/bin/env python3
"""
build_docx.py -- Convert CCE facilitator guide markdown files to formatted .docx

Usage:
    python build/build_docx.py                   # build all guides
    python build/build_docx.py guides/1sw/wk0-classroom-routines.md  # build one
"""

import sys, os, re, glob, yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Allow running from project root or build/
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "build"))
from config import *

GUIDES_DIR = ROOT / "cce-curriculum" / "guides"
OUTPUT_DIR = ROOT / "output" / "docx"


# ═══════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════

def set_font(run, name=FONT_NAME, size=SIZE_BODY, color=None, bold=None, italic=None):
    """Apply font properties to a run."""
    run.font.name = name
    run.font.size = size
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_shading(paragraph, fill_hex):
    """Add background shading to a paragraph."""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>')
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def add_bottom_border(paragraph, color_hex="2E75B6", size="3"):
    """Add a bottom border line under a paragraph."""
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:color="{color_hex}" w:sz="{size}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_page_margins(doc):
    """Set 0.75-inch margins on all sections."""
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN


# ═══════════════════════════════════════════════════════════════
# Markdown parsing
# ═══════════════════════════════════════════════════════════════

def parse_frontmatter(text):
    """Extract YAML frontmatter and body from markdown text."""
    m = re.match(r'^---\s*\n(.*?)\n---\s*\n', text, re.DOTALL)
    if not m:
        raise ValueError("Missing YAML frontmatter")
    meta = yaml.safe_load(m.group(1))
    body = text[m.end():]
    return meta, body


def parse_sections(body):
    """Split body into sections by ## headings. Returns list of (heading, content)."""
    parts = re.split(r'^## (.+)$', body, flags=re.MULTILINE)
    # parts[0] is text before first ##, then alternating heading/content
    sections = []
    for i in range(1, len(parts), 2):
        heading = parts[i].strip()
        content = parts[i + 1] if i + 1 < len(parts) else ""
        sections.append((heading, content.strip()))
    return sections


# ═══════════════════════════════════════════════════════════════
# Content rendering
# ═══════════════════════════════════════════════════════════════

def render_inline(paragraph, text, base_size=SIZE_BODY, base_color=None):
    """Render a line of text with **bold** and *italic* inline markers."""
    # Split on bold/italic markers
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size, color=base_color, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=base_size, color=base_color, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size, color=base_color)


def add_teacher_script(doc, text):
    """Add a teacher scripting paragraph: blue bold label, italic quote."""
    para = doc.add_paragraph()
    # Extract the quoted text after **Teacher:**
    m = re.match(r'\*\*Teacher:\*\*\s*"?(.*?)"?\s*$', text, re.DOTALL)
    label_text = "Teacher: "
    body_text = m.group(1) if m else text.replace("**Teacher:**", "").strip().strip('"')

    run = para.add_run(label_text)
    set_font(run, size=SIZE_BODY, color=BLUE, bold=True)
    run = para.add_run(f"\u201c{body_text}\u201d")
    set_font(run, size=SIZE_BODY, italic=True)
    return para


def add_flag_box(doc, tag, text, bg_hex, label_color):
    """Add a colored flag box paragraph."""
    para = doc.add_paragraph()
    add_shading(para, bg_hex)
    run = para.add_run(f"[{tag}] ")
    set_font(run, size=SIZE_FLAG, color=label_color, bold=True)
    run = para.add_run(text)
    set_font(run, size=SIZE_FLAG)
    return para


def add_warmup(doc, text):
    """Add a WARM-UP paragraph with yellow background."""
    para = doc.add_paragraph()
    add_shading(para, BG_YELLOW)
    run = para.add_run("WARM-UP: ")
    set_font(run, size=SIZE_BODY, color=ORANGE_TEXT, bold=True)
    run = para.add_run(text)
    set_font(run, size=SIZE_BODY)
    return para


def add_exit_ticket(doc, text):
    """Add an EXIT TICKET paragraph with green background."""
    para = doc.add_paragraph()
    add_shading(para, BG_GREEN)
    run = para.add_run("EXIT TICKET: ")
    set_font(run, size=SIZE_BODY, color=GREEN_TEXT, bold=True)
    run = para.add_run(text)
    set_font(run, size=SIZE_BODY)
    return para


def add_dok(doc, level, text):
    """Add a DOK question paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(f"DOK {level}: ")
    set_font(run, size=SIZE_BODY, color=RED_TEXT, bold=True)
    run = para.add_run(text)
    set_font(run, size=SIZE_BODY, italic=True)
    return para


def add_edp(doc, text):
    """Add an EDP step paragraph."""
    para = doc.add_paragraph()
    run = para.add_run("EDP: ")
    set_font(run, size=SIZE_BODY, color=PURPLE_EDP, bold=True, italic=True)
    run = para.add_run(text)
    set_font(run, size=SIZE_BODY, color=PURPLE_EDP, italic=True)
    return para


def add_section_heading(doc, title):
    """Add a ## section heading with blue text and bottom border."""
    para = doc.add_paragraph()
    run = para.add_run(title.upper())
    set_font(run, size=SIZE_SECTION_H2, color=BLUE, bold=True)
    add_bottom_border(para, "2E75B6", "3")
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_day_heading(doc, title):
    """Add a ### Day heading."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    set_font(run, size=SIZE_DAY_H3, color=NAVY, bold=True)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    return para


def add_activity_heading(doc, title):
    """Add a #### activity heading."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    set_font(run, size=SIZE_ACTIVITY_H4, color=NAVY, bold=True)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    return para


def add_sub_heading(doc, title):
    """Add a ### subsection heading (e.g., Differentiation subsections)."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    set_font(run, size=SIZE_ACTIVITY_H4, color=NAVY, bold=True)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    return para


def add_bullet(doc, text):
    """Add a bullet point with inline formatting."""
    para = doc.add_paragraph(style='List Bullet')
    render_inline(para, text, base_size=SIZE_BODY)
    return para


def add_body_paragraph(doc, text):
    """Add a normal body paragraph with inline formatting."""
    para = doc.add_paragraph()
    render_inline(para, text, base_size=SIZE_BODY)
    return para


def render_content_line(doc, line):
    """Render a single line of section content, detecting special markers."""
    stripped = line.strip()
    if not stripped:
        return

    # Blockquote lines
    if stripped.startswith("> "):
        inner = stripped[2:].strip()

        # Teacher script
        if inner.startswith("**Teacher:**"):
            add_teacher_script(doc, inner)
            return

        # H&L PLATFORM flag
        m = re.match(r'\[H&L PLATFORM\]\s*(.*)', inner)
        if m:
            add_flag_box(doc, "H&L PLATFORM", m.group(1), BG_LIGHT_BLUE, PURPLE_TEXT)
            return

        # VERIFY IN H&L flag
        m = re.match(r'\[VERIFY IN H&L\]\s*(.*)', inner)
        if m:
            add_flag_box(doc, "VERIFY IN H&L", m.group(1), BG_ORANGE, ORANGE_TEXT)
            return

        # VERIFY IN eDynamic flag
        m = re.match(r'\[VERIFY IN eDynamic\]\s*(.*)', inner)
        if m:
            add_flag_box(doc, "VERIFY IN eDynamic", m.group(1), BG_LIGHT_PURPLE, PURPLE_EDYN)
            return

        # "I can" statement (Demonstration of Learning)
        add_demonstration_of_learning(doc, inner)
        return

    # WARM-UP
    m = re.match(r'\*\*WARM-UP:\*\*\s*(.*)', stripped)
    if m:
        add_warmup(doc, m.group(1))
        return

    # EXIT TICKET
    m = re.match(r'\*\*EXIT TICKET:\*\*\s*(.*)', stripped)
    if m:
        add_exit_ticket(doc, m.group(1))
        return

    # DOK
    m = re.match(r'\*\*DOK (\d):\*\*\s*(.*)', stripped)
    if m:
        add_dok(doc, m.group(1), m.group(2))
        return

    # EDP
    m = re.match(r'\*\*EDP:\*\*\s*(.*)', stripped)
    if m:
        add_edp(doc, m.group(1))
        return

    # Bullet list item
    m = re.match(r'^- (.*)', stripped)
    if m:
        add_bullet(doc, m.group(1))
        return

    # Regular paragraph
    add_body_paragraph(doc, stripped)


def add_demonstration_of_learning(doc, text):
    """Add the 'I can...' statement with green shading."""
    para = doc.add_paragraph()
    add_shading(para, BG_GREEN)
    # Remove surrounding quotes if present
    clean = text.strip().strip('"').strip('\u201c\u201d')
    run = para.add_run(clean)
    set_font(run, size=SIZE_SUBTITLE, color=GREEN_TEXT, bold=True, italic=True)
    return para


def render_section_content(doc, heading, content):
    """Render a full ## section: heading + its content lines."""
    add_section_heading(doc, heading)

    # Inside Lesson Sequence, handle ### and #### subheadings
    is_lesson_seq = heading.lower().startswith("lesson sequence")
    is_differentiation = heading.lower().startswith("differentiation")

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ### subheading
        m = re.match(r'^### (.+)', stripped)
        if m:
            title = m.group(1).strip()
            if is_lesson_seq:
                add_day_heading(doc, title)
            else:
                add_sub_heading(doc, title)
            i += 1
            continue

        # #### subheading
        m = re.match(r'^#### (.+)', stripped)
        if m:
            add_activity_heading(doc, m.group(1).strip())
            i += 1
            continue

        # Multi-line blockquote: collect continuation lines
        if stripped.startswith("> "):
            block_lines = [stripped[2:].strip()]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("> "):
                i += 1
                block_lines.append(lines[i].strip()[2:].strip())
            full_inner = " ".join(block_lines)
            # Re-process as a single blockquote
            render_content_line(doc, "> " + full_inner)
            i += 1
            continue

        render_content_line(doc, line)
        i += 1


# ═══════════════════════════════════════════════════════════════
# Document assembly
# ═══════════════════════════════════════════════════════════════

def build_info_table(doc, meta):
    """Build the 2x2 info table below the title."""
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells = [
        ("Topic", meta.get("topic", "")),
        ("Length", meta.get("length", "")),
        ("H&L Cluster", meta.get("cluster", "")),
        ("Irving ISD Pathways", meta.get("pathways", "")),
    ]

    for idx, (label, value) in enumerate(cells):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)
        # Clear default paragraph
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(f"{label}: ")
        set_font(run, size=SIZE_BODY, color=BLUE, bold=True)
        run = cell.paragraphs[0].add_run(value)
        set_font(run, size=SIZE_BODY)

    # Light gray background for table cells
    for row in table.rows:
        for cell in row.cells:
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shd)

    return table


def build_docx(md_path, output_path):
    """Build a single .docx from a markdown facilitator guide."""
    text = Path(md_path).read_text(encoding="utf-8")
    meta, body = parse_frontmatter(text)
    sections = parse_sections(body)

    doc = Document()
    set_page_margins(doc)

    # ── Header block ──
    # "Career and College Explorations"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Career and College Explorations")
    set_font(run, size=SIZE_HEADER_TOP, color=GRAY_TITLE)

    # "Facilitator Guide"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Facilitator Guide")
    set_font(run, size=SIZE_HEADER_SUB, color=GRAY_TITLE)

    # "1st Six Weeks | Week 0"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{meta.get('six_weeks', '')} | {meta.get('week', '')}")
    set_font(run, size=SIZE_HEADER_SUB, color=GRAY_TITLE)

    # Spacer
    doc.add_paragraph()

    # ── Title ──
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(meta.get("title", ""))
    set_font(run, size=SIZE_TITLE, color=NAVY, bold=True)

    # ── Subtitle ──
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(meta.get("subtitle", ""))
    set_font(run, size=SIZE_SUBTITLE, color=GRAY_SUB, italic=True)

    # Spacer
    doc.add_paragraph()

    # ── Info table ──
    build_info_table(doc, meta)
    doc.add_paragraph()  # spacer after table

    # ── Sections ──
    for heading, content in sections:
        render_section_content(doc, heading, content)

    # ── Save ──
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════

def docx_filename_from_meta(meta):
    """Generate the output filename from frontmatter metadata."""
    sw = meta.get("six_weeks", "").split()[0]  # "1st" -> "1"
    sw_num = re.search(r'\d+', sw)
    sw_str = sw_num.group() if sw_num else "X"
    wk = meta.get("week", "Week 0").split()[-1]  # "Week 0" -> "0"
    title_slug = re.sub(r'[^a-zA-Z0-9]+', '_', meta.get("title", "guide")).strip('_')
    return f"FG_{sw_str}SW_Wk{wk}_{title_slug}.docx"


def main():
    if len(sys.argv) > 1:
        # Build specific file(s)
        md_files = [Path(f) for f in sys.argv[1:]]
    else:
        # Build all guides
        md_files = sorted(GUIDES_DIR.rglob("*.md"))

    if not md_files:
        print("No markdown files found.")
        return

    built = 0
    for md_path in md_files:
        try:
            text = md_path.read_text(encoding="utf-8")
            meta, _ = parse_frontmatter(text)
            out_name = docx_filename_from_meta(meta)
            out_path = OUTPUT_DIR / out_name
            build_docx(md_path, out_path)
            print(f"  OK  {md_path.name} -> {out_path.name}")
            built += 1
        except Exception as e:
            print(f"  FAIL  {md_path.name}: {e}")

    print(f"\nBuilt {built}/{len(md_files)} guides -> {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
